"""
AutoGrantED Core Engine
=======================
Zero-token grant discovery, orchestration, compliance & export.
Powered by SignalMesh 72-slot spatial frequency grid + SEC-Ω validation.
"""

from __future__ import annotations

import hashlib
import json
import os
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Set

from pydantic import BaseModel, Field


# =====================================================================
# 1. SIGNALMESH DETERMINISTIC HUB
# =====================================================================

class SignalEnvelope(BaseModel):
    uri: str
    slot: int
    frequencies: List[str]
    data: Dict[str, Any]
    timestamp: float = Field(default_factory=time.time)


class SpatialSignalMesh:
    """
    Deterministic 72-slot spatial routing (SHA-256(uri) % 72).
    Replaces vector DB lookups with microsecond in-memory hydration.
    """

    def __init__(self) -> None:
        self.grid: Dict[int, List[SignalEnvelope]] = {i: [] for i in range(72)}
        self.history: List[SignalEnvelope] = []

    def compute_slot(self, uri: str) -> int:
        return int(hashlib.sha256(uri.encode("utf-8")).hexdigest(), 16) % 72

    def emit(
        self,
        uri: str,
        payload: Dict[str, Any],
        frequencies: List[str],
    ) -> SignalEnvelope:
        slot = self.compute_slot(uri)
        envelope = SignalEnvelope(
            uri=uri,
            slot=slot,
            frequencies=frequencies,
            data=payload,
        )
        self.grid[slot].append(envelope)
        self.history.append(envelope)
        # Keep history bounded
        if len(self.history) > 500:
            self.history = self.history[-300:]
        return envelope

    def get_slot_signals(self, slot: int) -> List[SignalEnvelope]:
        return self.grid.get(slot, [])

    def get_grid_summary(self) -> Dict[str, Any]:
        populated = {
            str(slot): len(signals)
            for slot, signals in self.grid.items()
            if signals
        }
        return {
            "total_slots": 72,
            "populated_slots": populated,
            "total_signals": sum(len(s) for s in self.grid.values()),
            "recent_uris": [e.uri for e in self.history[-10:]],
        }

    def clear(self) -> None:
        self.grid = {i: [] for i in range(72)}
        self.history = []


# =====================================================================
# 2. SEC-Ω COMPLIANCE WARDEN
# =====================================================================

class SecOmegaEngine:
    """Deterministic validation gate for NSF PAPPG & DARPA BAA rules."""

    ALLOWED_NSF_FONTS = {
        "arial": 10.0,
        "times new roman": 11.0,
        "courier new": 10.0,
        "palatino linotype": 10.0,
        "computer modern": 11.0,
    }

    @classmethod
    def audit(cls, draft: Dict[str, Any]) -> Dict[str, Any]:
        agency = str(draft.get("agency", "NSF")).upper()
        formatting = draft.get("formatting", {}) or {}
        sections = draft.get("sections", {}) or {}
        errors: List[str] = []
        warnings: List[str] = []

        font = str(formatting.get("font_family", "Times New Roman")).lower()
        size = float(formatting.get("font_size", 11.0))
        margin = float(formatting.get("margin_inches", 1.0))

        if margin < 1.0:
            errors.append(f"Margin violation: {margin} in. Must be ≥ 1.0 in.")

        if agency == "NSF":
            if font not in cls.ALLOWED_NSF_FONTS:
                errors.append(f"Font '{font}' not permitted under NSF PAPPG.")
            elif size < cls.ALLOWED_NSF_FONTS[font]:
                errors.append(
                    f"Font size {size}pt below minimum "
                    f"{cls.ALLOWED_NSF_FONTS[font]}pt for {font}."
                )

            summary = sections.get("project_summary", {}) or {}
            for key in ("overview", "intellectual_merit", "broader_impacts"):
                if not summary.get(key) or not str(summary.get(key)).strip():
                    errors.append(
                        f"Project Summary missing mandatory subsection: '{key}'."
                    )

            desc = sections.get("project_description", {}) or {}
            pages = int(desc.get("estimated_pages", 0) or 0)
            if pages > 15:
                errors.append(
                    f"Project Description exceeds 15-page limit (found {pages})."
                )

            content = str(desc.get("content", ""))
            if "Broader Impacts" not in content and "broader impacts" not in content.lower():
                warnings.append(
                    "Project Description should contain an explicit "
                    "'Broader Impacts' header."
                )

            if "data_management_plan" not in sections:
                errors.append("Missing required document: Data Management Plan.")

        elif agency == "DARPA":
            if size < 12.0:
                errors.append(
                    f"DARPA BAA requires minimum 12pt font (received {size}pt)."
                )
            if "volume_1_technical" not in sections:
                errors.append("Missing mandatory Volume I (Technical Proposal).")
            admin = draft.get("admin_metadata", {}) or {}
            for field_name in ("baa_number", "proposer_org", "proposal_title", "uei"):
                if not admin.get(field_name):
                    errors.append(
                        f"Missing mandatory DARPA administrative field: '{field_name}'."
                    )

        else:
            # Generic fallback checks
            if not sections.get("project_summary"):
                warnings.append("No project_summary provided for non-NSF/DARPA agency.")

        return {
            "is_valid": len(errors) == 0,
            "agency": agency,
            "errors": errors,
            "warnings": warnings,
            "timestamp": time.time(),
        }


# =====================================================================
# 3. DOCUMENT COMPILER (LaTeX + DOCX)
# =====================================================================

class ProposalCompiler:
    EXPORT_DIR = os.environ.get("AUTOGRANTED_EXPORT_DIR", "./exports")

    @classmethod
    def ensure_dir(cls) -> None:
        os.makedirs(cls.EXPORT_DIR, exist_ok=True)

    @classmethod
    def _sanitize_tex(cls, text: str) -> str:
        if not text:
            return ""
        replacements = {
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    @classmethod
    def compile_latex(cls, submission_id: str, manifest: Dict[str, Any]) -> str:
        cls.ensure_dir()
        artifacts = manifest.get("compiled_artifacts", {}) or {}
        title = cls._sanitize_tex(
            artifacts.get("cover_sheet", {}).get("proposal_title", "Grant Proposal")
        )
        summary = artifacts.get("summary", {}) or {}
        narrative = cls._sanitize_tex(artifacts.get("core_narrative", "") or "")
        dmp = cls._sanitize_tex(
            str(artifacts.get("data_management_plan", "FAIR-compliant open repository."))
        )

        tex = f"""\\documentclass[11pt,letterpaper]{{article}}
\\usepackage[margin=1.0in]{{geometry}}
\\usepackage{{mathptmx}}
\\usepackage{{titlesec}}
\\usepackage{{hyperref}}
\\usepackage{{enumitem}}

\\titleformat{{\\section}}{{\\large\\bfseries}}{{\\thesection}}{{1em}}{{}}
\\titleformat{{\\subsection}}{{\\normalsize\\bfseries}}{{\\thesubsection}}{{1em}}{{}}

\\title{{{title}}}
\\author{{AutoGrantED Orchestration Engine}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle

\\section*{{Project Summary}}
\\subsection*{{Overview}}
{cls._sanitize_tex(summary.get("overview", "N/A"))}

\\subsection*{{Intellectual Merit}}
{cls._sanitize_tex(summary.get("intellectual_merit", "N/A"))}

\\subsection*{{Broader Impacts}}
{cls._sanitize_tex(summary.get("broader_impacts", "N/A"))}

\\newpage
\\section*{{Project Description}}
{narrative}

\\newpage
\\section*{{Data Management Plan}}
{dmp}

\\end{{document}}
"""
        filepath = os.path.join(cls.EXPORT_DIR, f"{submission_id}.tex")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(tex)
        return filepath

    @classmethod
    def compile_docx(cls, submission_id: str, manifest: Dict[str, Any]) -> Optional[str]:
        cls.ensure_dir()
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return None

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        artifacts = manifest.get("compiled_artifacts", {}) or {}
        title = artifacts.get("cover_sheet", {}).get(
            "proposal_title", "Grant Proposal"
        )

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        doc.add_heading("Project Summary", level=1)
        sum_data = artifacts.get("summary", {}) or {}
        for heading, key in [
            ("Overview", "overview"),
            ("Intellectual Merit", "intellectual_merit"),
            ("Broader Impacts", "broader_impacts"),
        ]:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(sum_data.get(key, "N/A"))

        doc.add_page_break()
        doc.add_heading("Project Description", level=1)
        doc.add_paragraph(artifacts.get("core_narrative", "") or "")

        doc.add_page_break()
        doc.add_heading("Data Management Plan", level=1)
        doc.add_paragraph(
            str(artifacts.get("data_management_plan", "FAIR-compliant open repository."))
        )

        filepath = os.path.join(cls.EXPORT_DIR, f"{submission_id}.docx")
        doc.save(filepath)
        return filepath


# =====================================================================
# 4. LIFECYCLE MONITOR
# =====================================================================

class GrantState(str, Enum):
    QUEUED = "QUEUED"
    SUBMITTED = "SUBMITTED"
    UNDER_REVIEW = "UNDER_REVIEW"
    REVISION_REQUESTED = "REVISION_REQUESTED"
    AWARDED = "AWARDED"
    REJECTED = "REJECTED"


class LifecycleMonitor:
    def __init__(self) -> None:
        self.submissions: Dict[str, Dict[str, Any]] = {}

    def track(
        self,
        sub_id: str,
        agency: str,
        target_latency_us: float = 1.69,
    ) -> None:
        self.submissions[sub_id] = {
            "agency": agency,
            "state": GrantState.SUBMITTED.value,
            "target_latency_us": target_latency_us,
            "created_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "last_update": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "milestones": [
                {"name": "Initial Compliance Check", "status": "COMPLETED"},
                {"name": "Panel Merit Review", "status": "PENDING"},
                {"name": "Program Officer Recommendation", "status": "PENDING"},
                {"name": "Award Verification", "status": "PENDING"},
            ],
            "notes": [],
        }

    def update_state(
        self, sub_id: str, new_state: GrantState, note: str = ""
    ) -> bool:
        if sub_id not in self.submissions:
            return False
        self.submissions[sub_id]["state"] = new_state.value
        self.submissions[sub_id]["last_update"] = time.strftime(
            "%Y-%m-%dT%H:%M:%SZ", time.gmtime()
        )
        if note:
            self.submissions[sub_id]["notes"].append(note)
        return True

    def verify_kpis(
        self, sub_id: str, observed_latency_us: float
    ) -> Dict[str, Any]:
        entry = self.submissions.get(sub_id)
        if not entry:
            return {"error": "Submission ID not found"}
        target = entry.get("target_latency_us", 1.69)
        return {
            "submission_id": sub_id,
            "state": entry["state"],
            "kpi_passed": observed_latency_us <= target,
            "observed_latency_us": observed_latency_us,
            "target_latency_us": target,
            "agency": entry["agency"],
        }

    def list_submissions(self) -> List[Dict[str, Any]]:
        return list(self.submissions.values())


# =====================================================================
# 5. GRANTS.GOV SCOUT
# =====================================================================

class GrantsGovScout:
    SEARCH_ENDPOINT = "https://api.grants.gov/v1/api/search2"

    def fetch(
        self,
        keywords: str = "smart grid AI orchestration",
        agencies: str = "NSF|DARPA|DOE",
        rows: int = 5,
    ) -> List[Dict[str, Any]]:
        payload = {
            "keyword": keywords,
            "agencies": agencies,
            "oppStatuses": "posted|forecasted",
            "rows": rows,
        }
        req = urllib.request.Request(
            self.SEARCH_ENDPOINT,
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "User-Agent": "AutoGrantED/2.0 (SignalMesh)",
            },
        )
        try:
            with urllib.request.urlopen(req, timeout=12) as response:
                res = json.loads(response.read().decode("utf-8"))
                return res.get("data", {}).get("oppHits", []) or []
        except Exception as exc:
            print(f"[GrantsGovScout] Live query failed: {exc}")
            return []

    def fallback_hits(self, keywords: str) -> List[Dict[str, Any]]:
        return [
            {
                "id": "NSF-2026-GRID-AI",
                "number": "NSF-26-501",
                "title": f"Decentralized AI Orchestration for Scalable Smart Energy Grids — {keywords}",
                "agencyCode": "NSF",
                "agencyName": "National Science Foundation",
                "openDate": "2026-08-01",
                "closeDate": "2026-11-30",
                "oppStatus": "posted",
            },
            {
                "id": "DARPA-AS-2026",
                "number": "HR001126S0001",
                "title": "Autonomous Multi-Agent Signal Mesh Coordination",
                "agencyCode": "DARPA",
                "agencyName": "Defense Advanced Research Projects Agency",
                "openDate": "2026-07-15",
                "closeDate": "2026-12-15",
                "oppStatus": "posted",
            },
        ]


# Global shared instances (used by app.py)
mesh = SpatialSignalMesh()
monitor = LifecycleMonitor()
scout = GrantsGovScout()
